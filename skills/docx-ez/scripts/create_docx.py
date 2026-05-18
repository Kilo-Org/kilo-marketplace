#!/usr/bin/env python3
"""
create_docx.py — Create a .docx document from scratch using python-docx.

Use this when Node.js is unavailable. For richer TOC / bookmark support
prefer create_docx.js (docx.js).

Usage:
  python create_docx.py <output.docx>                      # sample document
  python create_docx.py <output.docx> --spec <spec.json>   # from JSON spec

Spec JSON format (all fields optional):
{
  "title":    "Document Title",
  "author":   "Author Name",
  "subject":  "Subject",
  "sections": [
    {
      "children": [
        { "type": "heading",   "level": 1, "text": "Introduction" },
        { "type": "paragraph", "text": "Body text.", "bold": false, "italic": false },
        { "type": "table",
          "rows": [["Header A", "Header B"], ["r1c1", "r1c2"]] },
        { "type": "image",     "path": "image.png",
          "width_cm": 10, "height_cm": 6 },
        { "type": "page_break" }
      ]
    }
  ]
}

Dependency:
  pip install python-docx
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn
    from docx.shared import Cm, Inches, Pt, RGBColor
except ImportError:
    print('Error: python-docx not installed. Run: pip install python-docx',
          file=sys.stderr)
    sys.exit(1)


# ── Page / font defaults ──────────────────────────────────────────────────────

LETTER_W_CM  = 21.59   # 8.5 in
LETTER_H_CM  = 27.94   # 11 in
MARGIN_CM    = 2.54    # 1 in
DEFAULT_FONT = 'Arial'
DEFAULT_PT   = 12


def _apply_page_setup(doc: Document) -> None:
    """Set US Letter dimensions and 1-inch margins on all sections."""
    for section in doc.sections:
        section.page_width   = Cm(LETTER_W_CM)
        section.page_height  = Cm(LETTER_H_CM)
        section.top_margin    = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin   = Cm(MARGIN_CM)
        section.right_margin  = Cm(MARGIN_CM)


def _set_default_font(doc: Document) -> None:
    """Apply Arial 12pt as the document default run font."""
    style = doc.styles['Normal']
    font  = style.font
    font.name = DEFAULT_FONT
    font.size = Pt(DEFAULT_PT)
    # Also patch the theme font element so Word respects it
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(docx_qn('w:ascii'),    DEFAULT_FONT)
    rFonts.set(docx_qn('w:hAnsi'),    DEFAULT_FONT)
    rFonts.set(docx_qn('w:eastAsia'), DEFAULT_FONT)
    rFonts.set(docx_qn('w:cs'),       DEFAULT_FONT)
    existing = rPr.find(docx_qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


# ── Node builders ─────────────────────────────────────────────────────────────

HEADING_STYLES = {
    1: 'Heading 1',
    2: 'Heading 2',
    3: 'Heading 3',
    4: 'Heading 4',
    5: 'Heading 5',
    6: 'Heading 6',
}


def add_heading(doc: Document, level: int, text: str) -> None:
    style = HEADING_STYLES.get(level, 'Heading 1')
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = DEFAULT_FONT


def add_paragraph(doc: Document, node: dict) -> None:
    p    = doc.add_paragraph()
    text = node.get('text', '')
    run  = p.add_run(text)
    run.font.name  = DEFAULT_FONT
    run.font.size  = Pt(node.get('font_size', DEFAULT_PT))
    run.bold       = node.get('bold',   False)
    run.italic     = node.get('italic', False)
    if node.get('underline'):
        run.underline = True
    if node.get('color'):
        hex_color = node['color'].lstrip('#')
        r, g, b   = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    align = node.get('align', '').upper()
    if align == 'CENTER':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align in ('RIGHT', 'END'):
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align in ('JUSTIFY', 'BOTH'):
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_table(doc: Document, node: dict) -> None:
    rows_data = node.get('rows', [])
    if not rows_data:
        return
    ncols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows_data):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = str(cell_text)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else \
                  cell.paragraphs[0].add_run('')
            run.font.name = DEFAULT_FONT
            if ri == 0:  # header row bold
                run.bold = True


def add_image(doc: Document, node: dict) -> None:
    img_path = node.get('path', '')
    if not Path(img_path).exists():
        doc.add_paragraph(f'[Image not found: {img_path}]')
        return
    width_arg  = Cm(node['width_cm'])  if 'width_cm'  in node else Inches(3)
    height_arg = Cm(node['height_cm']) if 'height_cm' in node else None
    if height_arg:
        doc.add_picture(img_path, width=width_arg, height=height_arg)
    else:
        doc.add_picture(img_path, width=width_arg)


def add_toc_stub(doc: Document, title: str = 'Table of Contents') -> None:
    """Insert a TOC field placeholder. Requires field update in Word/LibreOffice."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = DEFAULT_FONT
    run.bold = True
    # Insert TOC field via raw XML
    fld_para = doc.add_paragraph()
    fld_para._p.clear()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(docx_qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(docx_qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(docx_qn('w:fldCharType'), 'end')
    r = OxmlElement('w:r')
    r.append(fldChar_begin)
    fld_para._p.append(r)
    r2 = OxmlElement('w:r')
    r2.append(instrText)
    fld_para._p.append(r2)
    r3 = OxmlElement('w:r')
    r3.append(fldChar_sep)
    fld_para._p.append(r3)
    r4 = OxmlElement('w:r')
    r4.append(fldChar_end)
    fld_para._p.append(r4)


def build_section(doc: Document, sec: dict) -> None:
    for node in sec.get('children', []):
        t = node.get('type', 'paragraph')
        if t == 'heading':
            add_heading(doc, node.get('level', 1), node.get('text', ''))
        elif t == 'paragraph':
            add_paragraph(doc, node)
        elif t == 'table':
            add_table(doc, node)
        elif t == 'image':
            add_image(doc, node)
        elif t == 'toc':
            add_toc_stub(doc, node.get('title', 'Table of Contents'))
        elif t == 'page_break':
            doc.add_page_break()
        else:
            doc.add_paragraph(str(node.get('text', '')))


# ── Document builder ───────────────────────────────────────────────────────────

def build_document(spec: dict) -> Document:
    doc = Document()
    _apply_page_setup(doc)
    _set_default_font(doc)

    # Core properties
    props = doc.core_properties
    if spec.get('title'):   props.title   = spec['title']
    if spec.get('author'):  props.author  = spec['author']
    if spec.get('subject'): props.subject = spec['subject']

    for sec in spec.get('sections', [{'children': []}]):
        build_section(doc, sec)

    return doc


# ── Sample spec ─────────────────────────────────────────────────────────────────

def sample_spec() -> dict:
    return {
        'title':  'Sample Document',
        'author': 'docx skill',
        'sections': [{
            'children': [
                {'type': 'heading', 'level': 1, 'text': 'Sample Document'},
                {'type': 'toc'},
                {'type': 'heading', 'level': 1, 'text': 'Introduction'},
                {'type': 'paragraph',
                 'text': 'This document was generated by the docx skill using '
                         'python-docx. It demonstrates headings, paragraphs, '
                         'tables, and a TOC field.'},
                {'type': 'heading', 'level': 2, 'text': 'Background'},
                {'type': 'paragraph',
                 'text': 'Key tools: python-docx for creation and editing, '
                         'lxml for direct XML manipulation, and LibreOffice '
                         'for headless PDF export.'},
                {'type': 'heading', 'level': 1, 'text': 'Toolchain'},
                {'type': 'table', 'rows': [
                    ['Task',            'Tool',           'Script'],
                    ['Create (Python)', 'python-docx',    'create_docx.py'],
                    ['Create (Node)',   'docx.js',        'create_docx.js'],
                    ['Edit existing',   'lxml',           'edit_docx.py'],
                    ['PDF export',      'LibreOffice',    'export_pdf.sh'],
                    ['Validate',        'Python zipfile', 'validate_docx.py'],
                ]},
                {'type': 'heading', 'level': 1, 'text': 'Formatting'},
                {'type': 'paragraph', 'text': 'Bold text example.',
                 'bold': True},
                {'type': 'paragraph', 'text': 'Italic text example.',
                 'italic': True},
                {'type': 'paragraph', 'text': 'Centred paragraph.',
                 'align': 'center'},
            ],
        }],
    }


# ── CLI ────────────────────────────────────────────────────────────────────────

def main() -> None:
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument('output', help='Output .docx path')
    ap.add_argument('--spec',  help='JSON spec file (optional)')
    args = ap.parse_args()

    if args.spec:
        try:
            spec = json.loads(Path(args.spec).read_text())
        except Exception as e:
            print(f'Error reading spec: {e}', file=sys.stderr)
            sys.exit(1)
    else:
        spec = sample_spec()

    doc = build_document(spec)
    doc.save(args.output)
    out = Path(args.output)
    print(f'Created: {args.output} ({out.stat().st_size:,} bytes)')


if __name__ == '__main__':
    main()
